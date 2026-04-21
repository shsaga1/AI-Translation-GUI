import argparse
import re
import sys
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from ebooklib import epub, ITEM_DOCUMENT
from pypdf import PdfReader

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from app.services.chunk_service import (
    normalize_text,
    split_text_into_blocks,
    split_markdown_into_blocks,
    build_context_chunks,
)
from app.services.translation_core import load_translation_model, translate_with_auto_review


def log(msg: str):
    print(msg, flush=True)


def status(msg: str):
    print(f"[STATUS] {msg}", flush=True)


def progress(val: int):
    print(f"[PROGRESS] {val}", flush=True)


def detect_file_type(input_file: str) -> str:
    ext = Path(input_file).suffix.lower()
    mapping = {
        ".md": "markdown",
        ".markdown": "markdown",
        ".docx": "docx",
        ".pdf": "pdf",
        ".epub": "epub",
    }
    return mapping.get(ext, "unknown")


def read_markdown_blocks(input_file: str) -> list[str]:
    text = Path(input_file).read_text(encoding="utf-8")
    return split_markdown_into_blocks(text)


def read_docx_blocks(input_file: str) -> list[str]:
    doc = Document(input_file)
    blocks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        text = normalize_text(text)
        if text:
            blocks.append(text)

    return blocks


def read_epub_blocks(input_file: str) -> list[str]:
    book = epub.read_epub(input_file)
    blocks = []

    for item in book.get_items():
        if item.get_type() != ITEM_DOCUMENT:
            continue

        soup = BeautifulSoup(item.get_content(), "html.parser")

        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        for node in soup.find_all(["h1", "h2", "h3", "p", "li", "blockquote"]):
            text = node.get_text(" ", strip=True)
            text = normalize_text(text)
            if len(text) >= 10:
                blocks.append(text)

    return blocks


def read_pdf_blocks(input_file: str) -> list[str]:
    reader = PdfReader(input_file)
    page_texts = []

    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.replace("\r", "\n").strip()

        if not text:
            continue

        text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        lines = [line.strip() for line in text.splitlines()]
        filtered_lines = []
        for line in lines:
            if not line:
                filtered_lines.append("")
                continue
            if len(line) <= 2 and line.isdigit():
                continue
            filtered_lines.append(line)

        page_text = "\n".join(filtered_lines).strip()
        if page_text:
            page_texts.append(page_text)

    merged_text = "\n\n".join(page_texts)
    merged_text = normalize_text(merged_text)

    return split_text_into_blocks(merged_text)


def read_blocks_by_type(file_type: str, input_file: str) -> list[str]:
    if file_type == "markdown":
        return read_markdown_blocks(input_file)
    if file_type == "docx":
        return read_docx_blocks(input_file)
    if file_type == "epub":
        return read_epub_blocks(input_file)
    if file_type == "pdf":
        return read_pdf_blocks(input_file)
    raise RuntimeError("未知文件类型，无法处理。")


def save_bilingual_markdown(
    output_dir: str,
    input_file: str,
    bilingual_blocks: list[tuple[str, str]],
    keep_original: bool,
) -> Path:
    input_path = Path(input_file)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    out_file = output_path / f"{input_path.stem}_bilingual.md"

    parts = [f"# 翻译结果：{input_path.name}", ""]

    for idx, (src, tgt) in enumerate(bilingual_blocks, start=1):
        parts.append(f"## 块 {idx}")
        parts.append("")

        if keep_original:
            parts.append("### 原文")
            parts.append("")
            parts.append(src)
            parts.append("")

        parts.append("### 译文")
        parts.append("")
        parts.append(tgt)
        parts.append("")

    out_file.write_text("\n".join(parts), encoding="utf-8")
    return out_file


def save_bilingual_docx(
    output_dir: str,
    input_file: str,
    bilingual_blocks: list[tuple[str, str]],
    keep_original: bool,
) -> Path:
    input_path = Path(input_file)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    out_file = output_path / f"{input_path.stem}_bilingual.docx"

    doc = Document()
    doc.add_heading(f"翻译结果：{input_path.name}", level=1)

    for idx, (src, tgt) in enumerate(bilingual_blocks, start=1):
        doc.add_heading(f"块 {idx}", level=2)

        if keep_original:
            doc.add_paragraph("原文：")
            doc.add_paragraph(src)

        doc.add_paragraph("译文：")
        doc.add_paragraph(tgt)

    doc.save(out_file)
    return out_file


def save_bilingual_html(
    output_dir: str,
    input_file: str,
    bilingual_blocks: list[tuple[str, str]],
    keep_original: bool,
) -> Path:
    input_path = Path(input_file)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    out_file = output_path / f"{input_path.stem}_bilingual.html"

    html_parts = [
        "<!DOCTYPE html>",
        "<html lang='zh'>",
        "<head>",
        "<meta charset='utf-8'>",
        "<meta name='viewport' content='width=device-width, initial-scale=1.0'>",
        f"<title>翻译结果：{escape_html(input_path.name)}</title>",
        "<style>",
        "body { font-family: Arial, sans-serif; max-width: 1100px; margin: 0 auto; padding: 24px; line-height: 1.7; }",
        ".block { border: 1px solid #ddd; border-radius: 10px; padding: 16px; margin-bottom: 18px; }",
        ".src { margin-bottom: 10px; color: #222; white-space: pre-wrap; }",
        ".tgt { color: #0b5394; background: #f5fbff; padding: 12px; border-radius: 8px; white-space: pre-wrap; }",
        "h1 { margin-bottom: 20px; }",
        "</style>",
        "</head>",
        "<body>",
        f"<h1>翻译结果：{escape_html(input_path.name)}</h1>",
    ]

    for idx, (src, tgt) in enumerate(bilingual_blocks, start=1):
        html_parts.append("<div class='block'>")
        html_parts.append(f"<h2>块 {idx}</h2>")

        if keep_original:
            html_parts.append(f"<div class='src'>{escape_html(src)}</div>")

        html_parts.append(f"<div class='tgt'>{escape_html(tgt)}</div>")
        html_parts.append("</div>")

    html_parts.extend(["</body>", "</html>"])
    out_file.write_text("\n".join(html_parts), encoding="utf-8")
    return out_file


def escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
    )


def choose_output_writer(file_type: str):
    if file_type == "markdown":
        return save_bilingual_markdown
    if file_type == "docx":
        return save_bilingual_docx
    if file_type == "epub":
        return save_bilingual_html
    if file_type == "pdf":
        return save_bilingual_html
    raise RuntimeError(f"当前文件类型暂不支持输出写回：{file_type}")


def translate_file(
    input_file: str,
    output_dir: str,
    mt_model_path: str,
    source_language: str,
    target_language: str,
    context_level_name: str,
    context_window: int,
    compute_precision: str,
    keep_original: bool,
    auto_review: bool,
):
    file_type = detect_file_type(input_file)

    status("校验输入")
    progress(5)
    log(f"[INFO] 输入文件: {input_file}")
    log(f"[INFO] 文件类型: {file_type}")
    log(f"[INFO] 输出目录: {output_dir}")
    log(f"[INFO] 翻译模型路径: {mt_model_path}")
    log(f"[INFO] 源语言: {source_language}")
    log(f"[INFO] 目标语言: {target_language}")
    log(f"[INFO] 上下文联系: {context_level_name}（窗口={context_window}）")
    log(f"[INFO] 推理精度: {compute_precision}")
    log(f"[INFO] 保留原文: {keep_original}")
    log(f"[INFO] 自动复查: {auto_review}")

    if file_type not in {"markdown", "docx", "epub", "pdf"}:
        raise RuntimeError("当前版本支持 Markdown / DOCX / EPUB / PDF。")

    status("读取文件中")
    progress(15)

    blocks = read_blocks_by_type(file_type, input_file)

    if not blocks:
        raise RuntimeError("未读取到可翻译内容。")

    log(f"[INFO] 读取到内容块数: {len(blocks)}")

    status("按上下文分块中")
    progress(25)

    chunks = build_context_chunks(blocks, context_window)
    log(f"[INFO] 当前上下文档位生成的翻译块数: {len(chunks)}")

    status("加载翻译模型")
    progress(35)
    tokenizer, model = load_translation_model(mt_model_path)

    bilingual_blocks = []
    total = len(chunks)

    for idx, chunk in enumerate(chunks, start=1):
        status(f"翻译中（{idx}/{total}）")
        current_progress = 35 + int((idx / max(total, 1)) * 55)
        progress(current_progress)

        src_text = "\n\n".join(chunk)
        log(f"[INFO] 正在翻译第 {idx}/{total} 块，字符数={len(src_text)}")

        translated = translate_with_auto_review(
            source_text=src_text,
            tokenizer=tokenizer,
            model=model,
            source_language=source_language,
            target_language=target_language,
            auto_review=auto_review,
            logger=log,
        )

        bilingual_blocks.append((src_text, translated))

    status("输出文件中")
    progress(95)

    writer = choose_output_writer(file_type)
    out_file = writer(
        output_dir=output_dir,
        input_file=input_file,
        bilingual_blocks=bilingual_blocks,
        keep_original=keep_original,
    )

    progress(100)
    status("完成")
    log(f"[INFO] 文件翻译完成，输出文件: {out_file}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input_file", required=True)
    parser.add_argument("--output_dir", required=True)
    parser.add_argument("--mt_model_path", required=True)
    parser.add_argument("--source_language", default="en")
    parser.add_argument("--target_language", default="zh")
    parser.add_argument("--context_level_name", default="标准上下文（中等）")
    parser.add_argument("--context_window", type=int, default=2)
    parser.add_argument("--compute_precision", default="auto")
    parser.add_argument("--keep_original", default="1")
    parser.add_argument("--auto_review", default="0")

    args = parser.parse_args()

    keep_original = str(args.keep_original) == "1"
    auto_review = str(args.auto_review) == "1"

    translate_file(
        input_file=args.input_file,
        output_dir=args.output_dir,
        mt_model_path=args.mt_model_path,
        source_language=args.source_language,
        target_language=args.target_language,
        context_level_name=args.context_level_name,
        context_window=args.context_window,
        compute_precision=args.compute_precision,
        keep_original=keep_original,
        auto_review=auto_review,
    )


if __name__ == "__main__":
    main()